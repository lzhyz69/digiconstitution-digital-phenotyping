from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "outputs" / "manuscript" / "plos_digital_health_full_manuscript_draft.md"
PLOS = ROOT / "outputs" / "plos_submission"
OUT_DIR = ROOT / "output" / "doc"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "plos_digital_health_full_manuscript_draft.docx"


TABLES = [
    (
        "Table 1. Baseline characteristics by constitution group.",
        PLOS / "tables" / "Table1_baseline_characteristics.csv",
        [
            "Continuous variables are shown as mean (standard deviation); binary variables are shown as number (percentage).",
            "P values compare all constitution groups using Kruskal-Wallis tests for continuous variables and chi-square tests for binary variables.",
            "SMD, standardized mean difference.",
        ],
    ),
    (
        "Table 2. Baseline constitution-to-disease and disease-marker associations.",
        PLOS / "tables" / "Table2_baseline_constitution_disease_associations.csv",
        [
            "Values are odds ratios with 95% confidence intervals.",
            "*False-discovery-rate-adjusted q<0.05.",
            "Balanced constitution is the reference group.",
        ],
    ),
    (
        "Table 3. Current constitution and next-visit disease-related or biochemical risk markers.",
        PLOS / "tables" / "Table3_lagged_next_visit_associations.csv",
        [
            "Values are odds ratios with 95% confidence intervals.",
            "*False-discovery-rate-adjusted q<0.05.",
            "Balanced constitution is the reference group.",
        ],
    ),
    (
        "Table 4. Temporal validation, calibration, and decision-curve metrics for digital constitution prediction.",
        PLOS / "tables" / "Table4_temporal_validation_prediction_metrics.csv",
        [
            "Training records were from 2017-2023, validation records from 2024, and temporal-test records from 2025-2026.",
            "Calibration and decision-curve metrics are shown for the full feature set.",
            "The sensitivity feature set excludes height, weight, BMI, and waist circumference.",
            "ECE, expected calibration error; PR-AUC, average precision.",
        ],
    ),
]


def get_section(text: str, start_heading: str, end_heading: str | None = None) -> str:
    start = text.index(start_heading)
    if end_heading is None:
        return text[start:].strip()
    end = text.index(end_heading, start + len(start_heading))
    return text[start:end].strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph_text(doc: Document, text: str, *, style: str | None = None, size: float = 10.5) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    format_run(run, size=size)


def add_markdown_block(doc: Document, block: str) -> None:
    for raw in block.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_paragraph_text(doc, line[2:], style="Title", size=16)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            add_paragraph_text(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_paragraph_text(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_paragraph_text(doc, line)


def autofit_table_font(table, font_size: float = 7.0) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    format_run(run, size=font_size)


def add_df_table(doc: Document, title: str, csv_path: Path, notes: list[str]) -> None:
    doc.add_heading(title, level=2)
    df = pd.read_csv(csv_path).fillna("")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, col in enumerate(df.columns):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E9EEF5")
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size=7.0, bold=True)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])

    autofit_table_font(table, font_size=7.0)
    for note in notes:
        add_paragraph_text(doc, f"Note: {note}", size=8.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)


def main() -> None:
    text = MANUSCRIPT.read_text(encoding="utf-8")
    body = text[: text.index("## Tables")].strip()
    figure_legends = get_section(text, "## Figure Legends", "## Supporting Information")
    supporting = get_section(text, "## Supporting Information", "## References")
    references = get_section(text, "## References")

    doc = Document()
    configure_document(doc)
    add_markdown_block(doc, body)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Tables", level=1)
    for title, csv_path, notes in TABLES:
        add_df_table(doc, title, csv_path, notes)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_block(doc, figure_legends)
    add_markdown_block(doc, supporting)
    add_markdown_block(doc, references)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
